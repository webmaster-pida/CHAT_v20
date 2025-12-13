# /src/main.py

import json
import asyncio
import io
import re
from datetime import datetime
from fastapi import FastAPI, Request, Depends, HTTPException, status, Form, Response
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any

# Librerías para documentos
from docx import Document
from fpdf import FPDF

from src.config import settings, log
from src.models.chat_models import ChatRequest, ChatMessage
from src.modules import vertex_search_client, gemini_client, rag_client, firestore_client
from src.core.prompts import PIDA_SYSTEM_PROMPT
from src.core.security import get_current_user

from google.cloud import firestore

app = FastAPI(
    title="PIDA Backend API",
    description="API para el asistente jurídico PIDA, con persistencia en BD y autenticación."
)

# --- CONFIGURACIÓN CORS ---
origins = settings.ALLOWED_ORIGINS

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_origin_regex=r"https://pida-ai-v20--.*\.web\.app$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- CLIENTE FIRESTORE ASÍNCRONO ---
db = firestore.AsyncClient()

# --- UTILIDADES DE NOMBRE DE ARCHIVO Y TEXTO (IGUAL QUE ANALIZADOR) ---
def generate_filename(title: str, extension: str) -> str:
    """Genera nombre seguro con fecha: Titulo_YYYY-MM-DD-HH-MM-SS.ext"""
    safe_title = re.sub(r'[^a-zA-Z0-9áéíóúÁÉÍÓÚñÑ ]', '', title[:40])
    safe_title = safe_title.strip().replace(' ', '_')
    if not safe_title:
        safe_title = "Chat_PIDA"
    timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
    return f"{safe_title}_{timestamp}.{extension}"

def sanitize_text_for_pdf(text: str) -> str:
    """Limpia caracteres incompatibles con Latin-1."""
    if not text: return ""
    replacements = {
        "•": "-", "—": "-", "–": "-", "“": '"', "”": '"', "‘": "'", "’": "'", "…": "...",
        "\u2013": "-", "\u2014": "-", "\u2022": "-", "\uF0B7": "-"
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text.encode('latin1', 'replace').decode('latin-1')

def write_markdown_to_pdf(pdf, text):
    """Interpreta Markdown básico para el PDF."""
    pdf.set_font("Arial", "", 11)
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue
            
        # Detectar roles de chat (ej: **Usuario**:)
        if line.startswith('**') and ':**' in line:
            parts = line.split(':**', 1)
            role = parts[0].replace('**', '')
            content = parts[1].strip()
            
            # Escribir Rol en Negrita y Color
            pdf.set_font("Arial", "B", 11)
            pdf.set_text_color(29, 53, 87) # Navy PIDA
            pdf.write(6, f"{role}: ")
            
            # Escribir contenido normal
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)
            
            # Procesar negritas internas en el contenido
            sub_parts = re.split(r'(\*\*.*?\*\*)', content)
            for sp in sub_parts:
                if sp.startswith('**') and sp.endswith('**'):
                    pdf.set_font("Arial", "B", 11)
                    pdf.write(6, sp.strip('*'))
                    pdf.set_font("Arial", "", 11)
                else:
                    pdf.write(6, sp)
            pdf.ln(6)
            
        # Títulos
        elif line.startswith('## '):
            pdf.ln(3)
            pdf.set_font("Arial", "B", 13)
            pdf.set_text_color(29, 53, 87)
            pdf.multi_cell(0, 8, line.replace('## ', ''))
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)
            
        # Listas
        elif line.startswith('* ') or line.startswith('- '):
            pdf.set_x(15)
            pdf.write(6, "- " + line[2:])
            pdf.ln(6)
            
        # Texto normal
        else:
            pdf.multi_cell(0, 6, line)

# --- CLASE PDF ---
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 14)
        self.set_text_color(29, 53, 87)
        self.cell(0, 10, "PIDA-AI: Historial de Chat", 0, 1, "L")
        self.set_font("Arial", "", 9)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Generado: {datetime.now().strftime('%d/%m/%Y, %H:%M:%S')}", 0, 1, "L")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Pagina {self.page_no()}/{{nb}}", 0, 0, "C")

# --- GENERADORES SÍNCRONOS ---
def create_chat_docx_sync(chat_text: str, title: str) -> tuple[bytes, str, str]:
    stream = io.BytesIO()
    doc = Document()
    doc.add_heading("PIDA-AI: Historial de Chat", 0)
    doc.add_paragraph(f"Tema: {title}")
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_heading("Conversación", 1)
    
    for line in chat_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
            
    doc.save(stream)
    stream.seek(0)
    fname = generate_filename(title, "docx")
    return stream.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", fname

def create_chat_pdf_sync(chat_text: str, title: str) -> tuple[bytes, str, str]:
    safe_text = sanitize_text_for_pdf(chat_text)
    safe_title = sanitize_text_for_pdf(title)
    
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # Título del Chat
    pdf.set_font("Arial", "B", 12)
    pdf.multi_cell(0, 6, f"Tema: {safe_title}")
    pdf.ln(5)
    
    # Contenido
    if not safe_text.strip():
        pdf.multi_cell(0, 6, "[Chat vacío]")
    else:
        write_markdown_to_pdf(pdf, safe_text)
        
    try:
        pdf_string = pdf.output(dest='S')
        if isinstance(pdf_string, str):
            pdf_bytes = pdf_string.encode('latin-1', 'replace')
        else:
            pdf_bytes = pdf_string
        stream = io.BytesIO(pdf_bytes)
        fname = generate_filename(title, "pdf")
        return stream.read(), "application/pdf", fname
    except Exception as e:
        print(f"Error PDF: {e}")
        err = FPDF()
        err.add_page()
        err.multi_cell(0, 10, f"Error: {str(e)}")
        return err.output(dest='S').encode('latin-1'), "application/pdf", "Error.pdf"

# --- VERIFICACIÓN DE SUSCRIPCIÓN ---
async def verify_active_subscription(current_user: Dict[str, Any]):
    user_id = current_user.get("uid")
    user_email = current_user.get("email", "").strip().lower()
    
    admin_domains = settings.ADMIN_DOMAINS
    admin_emails = settings.ADMIN_EMAILS
    email_domain = user_email.split("@")[-1] if "@" in user_email else ""

    if (email_domain in admin_domains) or (user_email in admin_emails):
        log.info(f"Acceso VIP concedido en Chat: {user_email}")
        return

    try:
        subscriptions_ref = db.collection("customers").document(user_id).collection("subscriptions")
        query = subscriptions_ref.where("status", "in", ["active", "trialing"]).limit(1)
        results = [doc async for doc in query.stream()]
        if not results:
            raise HTTPException(status_code=403, detail="No tienes una suscripción activa.")
    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        log.error(f"Error verificando suscripción DB: {e}")
        raise HTTPException(status_code=500, detail="Error interno verificando suscripción.")

# --- GENERADOR STREAMING ---
async def stream_chat_response_generator(chat_request: ChatRequest, country_code: str | None, user: Dict[str, Any], convo_id: str):
    user_id = user['uid']
    try:
        await verify_active_subscription(user) 
    except HTTPException as e:
        yield f"data: {json.dumps({'error': e.detail})}\n\n"
        return
    
    def create_sse_event(data: dict) -> str:
        return f"data: {json.dumps(data)}\n\n"

    try:
        user_message = ChatMessage(role="user", content=chat_request.prompt)
        await firestore_client.add_message_to_conversation(user_id, convo_id, user_message)
        
        yield create_sse_event({"event": "status", "message": "Iniciando... 🕵️"})
        await asyncio.sleep(0.1) 
        
        history_from_db = await firestore_client.get_conversation_messages(user_id, convo_id)
        history_for_gemini = gemini_client.prepare_history_for_vertex(history_from_db[:-1])
        
        yield create_sse_event({"event": "status", "message": "Consultando jurisprudencia..."})
        
        search_tasks = [
            asyncio.to_thread(vertex_search_client.search, chat_request.prompt, num_results=3),
            rag_client.search_internal_documents(chat_request.prompt)
        ]
        
        combined_context = ""
        for i, task in enumerate(asyncio.as_completed(search_tasks)):
            result = await task
            combined_context += result
            yield create_sse_event({"event": "status", "message": f"Fuente {i+1} procesada..."})
        
        yield create_sse_event({"event": "status", "message": "Generando respuesta..."})
        
        final_prompt = f"Contexto geográfico: {country_code}\n{combined_context}\n\n---\n\nPregunta del usuario: {chat_request.prompt}"
        
        full_response_text = ""
        async for chunk in gemini_client.generate_streaming_response(
            system_prompt=PIDA_SYSTEM_PROMPT,
            prompt=final_prompt,
            history=history_for_gemini
        ):
            yield create_sse_event({'text': chunk})
            full_response_text += chunk

        if full_response_text:
            model_message = ChatMessage(role="model", content=full_response_text)
            await firestore_client.add_message_to_conversation(user_id, convo_id, model_message)

        yield create_sse_event({'event': 'done'})

    except Exception as e:
        log.error(f"Error crítico streaming convo {convo_id}: {e}", exc_info=True)
        error_message = json.dumps({"error": "Ocurrió un error interno al generar la respuesta."})
        yield f"data: {error_message}\n\n"

# --- ENDPOINTS ---

@app.get("/status", tags=["Status"])
def read_status():
    return {"status": "ok", "message": "PIDA Chat Backend v3.0 (PDF Fixed)"}

@app.get("/conversations", response_model=List[Dict[str, Any]], tags=["Chat History"])
async def get_user_conversations(current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    return await firestore_client.get_conversations(current_user['uid'])

@app.get("/conversations/{convo_id}/messages", response_model=List[ChatMessage], tags=["Chat History"])
async def get_conversation_details(convo_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    return await firestore_client.get_conversation_messages(current_user['uid'], convo_id)

@app.post("/conversations", response_model=Dict[str, Any], status_code=status.HTTP_201_CREATED, tags=["Chat History"])
async def create_new_empty_conversation(request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    body = await request.json()
    title = body.get("title", "Nuevo Chat")
    if not title: raise HTTPException(400, "El título no puede estar vacío")
    new_convo = await firestore_client.create_new_conversation(current_user['uid'], title)
    return new_convo

@app.delete("/conversations/{convo_id}", status_code=status.HTTP_204_NO_CONTENT, tags=["Chat History"])
async def delete_a_conversation(convo_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    await firestore_client.delete_conversation(current_user['uid'], convo_id)
    return

@app.patch("/conversations/{convo_id}/title", status_code=status.HTTP_204_NO_CONTENT, tags=["Chat History"])
async def update_conversation_title_handler(convo_id: str, request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    body = await request.json()
    new_title = body.get("title")
    if not new_title: raise HTTPException(400, "El título no puede estar vacío")
    await firestore_client.update_conversation_title(current_user['uid'], convo_id, new_title)
    return

@app.post("/chat-stream/{convo_id}", tags=["Chat"])
async def chat_stream_handler(convo_id: str, chat_request: ChatRequest, request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    country_code = request.headers.get('X-Country-Code', None)
    headers = { "Content-Type": "text/event-stream", "Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no" }
    return StreamingResponse(stream_chat_response_generator(chat_request, country_code, current_user, convo_id), headers=headers)

# --- NUEVO ENDPOINT DE DESCARGA (SOLUCIÓN PDF/DOCX) ---
@app.post("/download-chat", tags=["Chat"])
async def download_chat(
    chat_text: str = Form(...),
    title: str = Form(...),
    file_format: str = Form("docx"),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Genera un archivo PDF o DOCX a partir del texto del chat.
    El formato de texto esperado es:
    **Rol**: Mensaje
    ...
    """
    try:
        if file_format.lower() == "docx":
            content, mime, fname = await asyncio.to_thread(create_chat_docx_sync, chat_text, title)
        else:
            content, mime, fname = await asyncio.to_thread(create_chat_pdf_sync, chat_text, title)
            
        return Response(content=content, media_type=mime, headers={"Content-Disposition": f"attachment; filename={fname}"})
    except Exception as e:
        log.error(f"Error descarga chat: {e}")
        raise HTTPException(500, f"Error generando archivo: {e}")

# --- NUEVO ENDPOINT DE VERIFICACIÓN VIP PARA EL FRONTEND ---
@app.post("/check-vip-access", tags=["Security"])
async def check_vip_access_handler(current_user: Dict[str, Any] = Depends(get_current_user)):
    """
    Endpoint simple para que el frontend verifique si el usuario tiene acceso VIP 
    antes de la verificación de suscripción de Stripe.
    
    NOTA: La lógica de acceso VIP (ADMIN_DOMAINS/EMAILS) se aplica dentro de get_current_user.
    Si el usuario llega aquí, significa que get_current_user NO le denegó el acceso.
    Por lo tanto, si get_current_user NO levantó 403, significa que el usuario es VIP.
    """
    
    # 1. Verificar Stripe (Lo hacemos aquí también, para ser coherentes con verify_active_subscription)
    user_id = current_user.get("uid")
    user_email = current_user.get("email", "").strip().lower()
    
    # La lógica de seguridad ya verificó si es VIP. 
    # Si llegó aquí, es un usuario válido. Ahora, ¿tiene suscripción?
    try:
        subscriptions_ref = db.collection("customers").document(user_id).collection("subscriptions")
        query = subscriptions_ref.where("status", "in", ["active", "trialing"]).limit(1)
        results = [doc async for doc in query.stream()]
        
        # Si tiene suscripción, el acceso está garantizado.
        if results:
            return jsonify({"is_vip_user": True})
            
    except Exception as e:
        log.error(f"Error verificando suscripción DB: {e}")
        # Si falla la DB, por seguridad, devolvemos False, pero es un fallo interno.
        # Dejamos que la lógica de abajo determine el acceso.
        pass 
        
    # Si la suscripción no está activa, chequeamos las reglas VIP/ADMIN.
    # Como get_current_user ya aplicó la restricción, si el usuario aún tiene acceso 
    # a este endpoint, es porque cumplió con las reglas de seguridad
    # (es decir, es un admin/VIP).

    # La única forma en que get_current_user permite el paso es si:
    # 1) No hay restricciones (has_restrictions es False).
    # 2) O si el email/dominio está en la lista (is_domain/email_authorized es True).
    
    # Dado que el frontend solo llama a esto si *falló* el chequeo de Stripe,
    # solo queremos saber si el usuario pasó la seguridad por ser VIP/Admin.
    
    # Si el usuario llega a esta línea, es porque get_current_user lo dejó pasar.
    # Si get_current_user lo dejó pasar Y no tiene Stripe, debe ser porque es VIP.
    
    admin_domains = settings.ADMIN_DOMAINS
    admin_emails = settings.ADMIN_EMAILS
    email_domain = user_email.split("@")[-1] if "@" in user_email else ""

    if (email_domain in admin_domains) or (user_email in admin_emails):
        log.info(f"Acceso VIP/Admin confirmado para {user_email}.")
        return {"is_vip_user": True}
        
    # Si no tiene suscripción activa y no es VIP, el acceso debe ser False.
    return {"is_vip_user": False}
